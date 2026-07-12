"""
dump_all_to_word.py
-------------------
Finds every .py file in the same directory as THIS script,
executes each one, and appends the result to a single Word document.

Output format per file:
    ──────────────────────
    filename.py
    ──────────────────────
    Code:
        <source code>

    Output:
        <captured stdout / stderr>

Dependencies:
    pip install python-docx
"""

import os
import sys
import subprocess
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Helpers ───────────────────────────────────────────────────────────────────

def run_script(path: str, timeout: int = 30) -> str:
    """Run a Python file and return its stdout + stderr as a string."""
    try:
        result = subprocess.run(
            [sys.executable, path],
            capture_output=True,
            text=True,
            timeout=timeout,
            cwd=os.path.dirname(path),
        )
        out = result.stdout
        if result.stderr:
            out += ("\n[stderr]\n" if out else "[stderr]\n") + result.stderr
        return out.rstrip("\n") or "(no output)"
    except subprocess.TimeoutExpired:
        return f"[ERROR] Timed out after {timeout}s."
    except Exception as e:
        return f"[ERROR] {e}"


def shade_para(para, hex_fill: str = "F2F2F2"):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_fill)
    pPr.append(shd)


def add_label(doc, text: str, color=(30, 30, 30)):
    """Bold label like 'Code:' or 'Output:'."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name      = "Calibri"
    run.font.size      = Pt(11)
    run.font.bold      = True
    run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)


def add_code_block(doc, text: str, fill: str = "F2F2F2"):
    """Render monospaced shaded lines."""
    for line in text.splitlines():
        p = doc.add_paragraph()
        shade_para(p, fill)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9.5)


def add_divider(doc, filename: str):
    """Section header with filename."""
    # Horizontal rule via border on the paragraph below
    p = doc.add_paragraph()
    run = p.add_run(filename)
    run.font.name      = "Calibri"
    run.font.size      = Pt(13)
    run.font.bold      = True
    run.font.color.rgb = RGBColor(0, 70, 127)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(2)


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    # Directory where the script is EXECUTED FROM (cwd)
    here = os.getcwd()
    this_file = os.path.abspath(__file__)

    # Collect all .py files in cwd (not subdirs, not this script itself)
    py_files = sorted(
        os.path.join(here, f)
        for f in os.listdir(here)
        if f.endswith(".py") and os.path.abspath(os.path.join(here, f)) != this_file
    )

    if not py_files:
        print("No .py files found in", here)
        return

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    print(f"Found {len(py_files)} file(s) in '{here}'\n")

    for path in py_files:
        fname = os.path.basename(path)
        print(f"  Processing: {fname} ...", end=" ", flush=True)

        # Read source
        try:
            with open(path, "r", encoding="utf-8") as f:
                code = f.read()
        except Exception as e:
            code = f"[Could not read file: {e}]"

        # Run and capture output
        output = run_script(path)
        print("done")

        # ── Write to doc ──
        add_divider(doc, fname)

        add_label(doc, "Code:")
        add_code_block(doc, code, fill="F2F2F2")

        add_label(doc, "Output:")
        add_code_block(doc, output, fill="FAFAFA")

        # Small gap between files
        doc.add_paragraph()

    out_path = os.path.join(here, "all_programs_output.docx")
    doc.save(out_path)
    print(f"\nSaved → {out_path}")
    print(f"(Scanned directory: {here})")


if __name__ == "__main__":
    main()
