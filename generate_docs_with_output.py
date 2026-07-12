"""
generate_docs_with_output.py
────────────────────────────
Scans the current working directory (recursively) for Python files whose
names match the pattern  P<number>_<Name>.py  (e.g. P01_1D_Array.py,
P204_Group_Stats.py), executes each one, captures its stdout/stderr, and
produces one Word (.docx) document per file.

Document layout
───────────────
Program <number>: <Name with underscores replaced by spaces>

Code:
    <source code of the program>

Output:
    <captured output from running the program>

Dependencies:  python-docx
    pip install python-docx
"""

import os
import re
import sys
import subprocess
import textwrap
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── helpers ───────────────────────────────────────────────────────────────────

def parse_filename(stem: str):
    """
    Extract (program_number, program_name) from a stem like 'P01_1D_Array'
    or 'P204_Group_Stats'.

    Returns (number_str, name_str) where name_str has underscores → spaces.
    Returns (None, None) if the stem does not match the expected pattern.
    """
    match = re.match(r"^[Pp](\d+)_(.+)$", stem)
    if not match:
        return None, None
    number = match.group(1).lstrip("0") or "0"   # '01' → '1'
    name   = match.group(2).replace("_", " ")
    return number, name


def run_program(py_path: str, timeout: int = 30) -> str:
    """
    Execute the given Python file using the current interpreter.
    Returns its combined stdout + stderr as a single string.
    Caps execution time at `timeout` seconds.
    """
    try:
        result = subprocess.run(
            [sys.executable, py_path],
            capture_output=True,
            text=True,
            timeout=timeout,
            cwd=os.path.dirname(py_path),   # run from the file's own directory
        )
        output = result.stdout
        if result.stderr:
            output += ("\n[stderr]\n" if output else "[stderr]\n") + result.stderr
        return output.rstrip("\n") or "(no output)"
    except subprocess.TimeoutExpired:
        return f"[ERROR] Execution timed out after {timeout} seconds."
    except Exception as exc:
        return f"[ERROR] Could not run program: {exc}"


def set_font(run, name="Courier New", size=10, bold=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_run(para, text, size=13, bold=True, color=(30, 30, 30)):
    run = para.add_run(text)
    run.font.name       = "Calibri"
    run.font.size       = Pt(size)
    run.font.bold       = bold
    run.font.color.rgb  = RGBColor(*color)


def shade_paragraph(para, fill_hex="F2F2F2"):
    """Apply a light grey background to a paragraph (simulates a code block)."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    pPr.append(shd)


def add_code_block(doc, text: str, fill_hex="F2F2F2"):
    """Render each line of `text` as a shaded monospaced paragraph."""
    for line in text.splitlines():
        p = doc.add_paragraph()
        shade_paragraph(p, fill_hex)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line if line else " ")   # keep blank lines visible
        set_font(run, name="Courier New", size=9.5)


# ── main logic ────────────────────────────────────────────────────────────────

def create_document(py_path: str, out_dir: str):
    stem = os.path.splitext(os.path.basename(py_path))[0]
    number, name = parse_filename(stem)

    if number is None:
        print(f"  [SKIP] '{stem}' – doesn't match P<num>_<name> pattern")
        return

    # ── Read source code ──
    with open(py_path, "r", encoding="utf-8") as f:
        code = f.read()

    # ── Execute program and capture output ──
    print(f"  [RUN]  Executing '{os.path.basename(py_path)}' …", end=" ", flush=True)
    output = run_program(py_path)
    print("done")

    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Title: "Program <n>: <Name>" ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_heading_run(
        title_para,
        f"Program {number}: {name}",
        size=14,
        bold=True,
        color=(0, 70, 127),          # dark blue
    )
    doc.add_paragraph()              # spacer

    # ── "Code:" label ──
    code_label = doc.add_paragraph()
    add_heading_run(code_label, "Code:", size=12, bold=True, color=(30, 30, 30))

    # ── Code block ──
    add_code_block(doc, code, fill_hex="F2F2F2")

    doc.add_paragraph()              # spacer

    # ── "Output:" label ──
    out_label = doc.add_paragraph()
    add_heading_run(out_label, "Output:", size=12, bold=True, color=(30, 30, 30))

    # ── Actual program output ──
    add_code_block(doc, output, fill_hex="FAFAFA")

    # ── Save ──
    out_path = os.path.join(out_dir, f"{stem}.docx")
    doc.save(out_path)
    print(f"  [OK]   '{stem}.docx' saved")


def find_py_files(root: str, script_name: str):
    """
    Walk `root` recursively and yield absolute paths to all .py files that
    match the P<num>_<name>.py pattern, excluding this script itself.
    """
    script_abs = os.path.abspath(script_name)
    for dirpath, _dirs, files in os.walk(root):
        for fname in sorted(files):
            if not fname.endswith(".py"):
                continue
            full = os.path.abspath(os.path.join(dirpath, fname))
            if full == script_abs:
                continue
            stem = os.path.splitext(fname)[0]
            number, _ = parse_filename(stem)
            if number is not None:
                yield full


def main():
    cwd = os.getcwd()
    out_dir = os.path.join(cwd, "Word_Documents")
    os.makedirs(out_dir, exist_ok=True)

    py_files = list(find_py_files(cwd, __file__))

    if not py_files:
        print("No matching Python files found in", cwd)
        return

    print(f"Found {len(py_files)} Python file(s) – output → '{out_dir}'\n")
    for fpath in py_files:
        # Show relative path for readability
        rel = os.path.relpath(fpath, cwd)
        print(f"→ {rel}")
        create_document(fpath, out_dir)
        print()

    print("Done!")


if __name__ == "__main__":
    main()
